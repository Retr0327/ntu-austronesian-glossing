from __future__ import annotations

from dataclasses import dataclass
from typing import (
    TYPE_CHECKING,
    Iterable,
)

import docx

from .base import (
    GlossReader,
    GlossWriter,
)

if TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document

    from glosstools.typings import FileDescriptorOrPath


@dataclass
class DocxGlossReader(GlossReader):
    file_path: FileDescriptorOrPath
    encoding: str = "utf-8"

    def __enter__(self):
        self.__file: Document = docx.Document(self.file_path)
        self.__content = list(
            map(lambda paragraph: paragraph.text, self.__file.paragraphs)
        )
        return self

    def __exit__(self, *args, **kwargs):
        return None

    def read(self) -> list[str]:
        content_without_newlines = map(self._strip_line, self.__content)
        return self._renumber(list(content_without_newlines))


@dataclass
class DocxGlossWriter(GlossWriter):
    file_path: Path
    data: Iterable[str]
    encoding: str = "utf-8"

    def __enter__(self):
        has_path = self.file_path.exists()
        if not has_path:
            self._mkdir(self.file_path)

        self.__doc: Document = docx.Document()
        for line in self.data:
            self.__doc.add_paragraph(line)
        return self

    def __exit__(self, *args, **kwargs):
        return None

    def write(self):
        return self.__doc.save(self.file_path)
